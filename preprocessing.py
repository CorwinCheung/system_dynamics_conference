import os
import pandas as pd
from docx import Document
import numpy as np
import re

cur_path = os.path.dirname(__file__)
true_path = "/Users/corwincheung/Programming/Systems Dynamics Analysis/data"


def explode_for_multiselect(df, column, specified_options):
    # Create a copy of the DataFrame
    df_copy = df.copy()

    # Split the column's values
    df_copy[column] = df_copy[column].str.split(', ')

    # Explode the DataFrame on the column
    df_copy = df_copy.explode(column)

    # Filter rows based on column's membership in the 'specified_options' list
    df_copy = df_copy[df_copy[column].isin(specified_options)]

    # Get value counts as a dictionary
    value_counts_dict = df_copy[column].value_counts().to_dict()

    return value_counts_dict


def year_2013():
    dfs_2013 = []
    document_2013 = Document("data/2013 Report survey.docx")
    # print out all columns for categorization

    # for paragraph in document_2013.paragraphs:
    #     if any(run.bold for run in paragraph.runs):
    #         if not any(substring in paragraph.text for substring in ["Text Entry", "View More", "%", "Answer"]):
    #             print(paragraph.text)

    for table in document_2013.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2013.append(df)

    print(len(dfs_2013))
    print(type(dfs_2013))

    print(type(dfs_2013[0]))

    print("data_type: " + str(type(dfs_2013)))
    print("data_type item: " + str(type(dfs_2013[0])))
    # print(dfs_2013[0:7])

    # categorical dictionaries

    # manual counting for ugly text data

    print(dfs_2013[38])

    slice_3 = dfs_2013[36][3]
    slice_3b = dfs_2013[36][1]
    slice_3b = [element.strip() for element in slice_3b]

    value_counts_dict_3 = {
        slice_3b[1]: int(slice_3[1]), slice_3b[2]: int(slice_3[2]), slice_3b[3]: int(slice_3[3]), slice_3b[4]: int(slice_3[4]),
        slice_3b[3]: int(slice_3[3]), slice_3b[6]: int(slice_3[6]), slice_3b[3]: int(slice_3[3]), slice_3b[8]: int(slice_3[8])
    }

    slice_4 = dfs_2013[38][3]
    slice_4b = dfs_2013[38][1]
    slice_4b = [element.strip() for element in slice_4b]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3]), slice_4b[4]: int(slice_4[4]),
        slice_4b[5]: int(slice_4[5]), slice_4b[6]: int(slice_4[6]), slice_4b[4]: int(slice_4[4]), slice_4b[8]: int(slice_4[8]),
        slice_4b[9]: int(slice_4[9]), slice_4b[10]: int(slice_4[10]), slice_4b[11]: int(slice_4[11]), slice_4b[12]: int(slice_4[12]),
        slice_4b[13]: int(slice_4[13]), slice_4b[14]: int(slice_4[14])
    }

    slice_5 = dfs_2013[41][3]
    slice_5b = dfs_2013[41][1]
    slice_5b = [element.strip() for element in slice_5b]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[5]: int(slice_5[5]),
        slice_5b[5]: int(slice_5[5]), slice_5b[6]: int(slice_5[6]), slice_5b[7]: int(slice_5[7])
    }

    value_counts_dict_6 = {
        'More than 20 years': 1, '10 to 20 years': 2, '2 years': 1, '3 years': 1,
        '1 year or less': 3, '5 to 10 years': 0, '4 years': 1
    }

    slice_2 = dfs_2013[39][3]
    slice_2b = dfs_2013[39][1]
    slice_2b = [element.strip() for element in slice_2b]

    value_counts_dict_2 = {
        slice_2b[1]: int(slice_2[1]), slice_2b[2]: int(slice_2[2]), slice_2b[3]: int(slice_2[3]), slice_2b[4]: int(slice_2[4])
    }

    slice_1 = dfs_2013[47][3]
    slice_1b = dfs_2013[47][1]
    slice_1b = [element.strip() for element in slice_1b]

    value_counts_dict_1 = {
        slice_1b[1]: int(slice_1[1]), slice_1b[2]: int(slice_1[2]), slice_1b[3]: int(slice_1[3]), slice_1b[4]: int(slice_1[4]), slice_1b[5]: int(slice_1[5])
    }

    # print(value_counts_dict_1)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6 = dicts

    value_counts_dicts = {
        "Have you attended the SD conference before?": value_counts_dict_2,
        "What is your profession? Please select all that apply.": value_counts_dict_3,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_4,
        "What is your geographic region?": value_counts_dict_5,
        "If you have used the Conference Online Scheduler, which device have you used? Please check all that apply.": value_counts_dict_1,
        "How many years of experience do you have with system dynamics?": value_counts_dict_6,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2013' not in df.index:
        df.loc['2013'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2013', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2013', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # boolean dictionaries

    boolean_value_counts_dicts = {
        "An idea that is considered for future conferences, is to make the conference schedule available using a smartphone app. For that reason, we would like to ask you if you own a smartphone.": {"Yes": 110, "No": 22},
        "Have you used the Conference Online Scheduler?": {"Yes": 45, "No": 86},
    }

    print(boolean_value_counts_dicts)

    df = pd.read_csv("boolean_questions.csv", index_col=0)

    if '2013' not in df.index:
        df.loc['2013'] = pd.NA

    for col_name, col_data in boolean_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2013', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2013', col_name] = col_data

    df.to_csv("boolean_questions.csv")

    # numerical dictionaries

    print(dfs_2013[20])

    numerical_indices = [1, 3, 5, 7, 9, 11, 13, 15, 17, 20, 21, 23,
                         24, 25, 28, 29, 30]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        print(idx)
        slice_main = dfs_2013[idx][3]
        slice_helper = dfs_2013[idx][1]
        slice_helper = [element.strip() for element in slice_helper]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to conference geographical location my evaluation is:",
        "When it comes to the conference venue (building and facilities) my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "When it comes to the conference fee my evaluation is:",
        "When it comes to the overnight accommodation my evaluation is:",
        "When it comes to the printed abstract proceedings for use at the conference my evaluation is:",
        "Are the plenary sessions too short or too long?",
        "Should we have more or fewer plenary sessions?",
        "Are the parallel sessions too short or too long?",
        "Should we have more or fewer parallel sessions?",
        "How do you evaluate the following sessions and workshops? [In-presence poster session]",
        "What is your opinion about the duration of the breaks during the conference?",
        "Would you like to see more or less academic work (as compared to applications)?",
        "How satisfied are you with the workshops?",
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2013' not in df.index:
        df.loc['2013'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2013', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2013', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    # print(dfs_2013[26])

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "Additional comments 7",
        "Additional comments 8",
        "If you have any comments on session formats, please list them here.",
        "If you could make changes to the format of the workshops, what would those changes be?",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify)",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "Any other comments you may have"
    ]
    text_list = []
    text_indices = [2, 4, 6, 8, 10, 12, 14, 16, 18,
                    26, 31, 32, 33, 34, 37, 40, 42, 48]

    for i in text_indices:
        text_list.append(dfs_2013[i][0][1:].values)

    text_value_counts_dicts = dict(zip(questions, text_list))
    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2013' not in df.index:
        df.loc['2013'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2013', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2013', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2014():
    dfs_2014 = []
    document_2014 = Document("data/2014 System dynamics conference.docx")
    # print out all columns for categorization

    # for paragraph in document_2014.paragraphs:
    #     if any(run.bold for run in paragraph.runs):
    #         if not any(substring in paragraph.text for substring in ["Text Entry", "View More", "%", "Answer"]):
    #             print(paragraph.text)

    for table in document_2014.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2014.append(df)

    print(len(dfs_2014))
    print(type(dfs_2014))

    print(type(dfs_2014[0]))

    print("data_type: " + str(type(dfs_2014)))
    print("data_type item: " + str(type(dfs_2014[0])))
    # print(dfs_2014[0:7])

    # categorical dictionaries

    # manual counting for ugly text data

    print(dfs_2014[35])

    slice_3 = dfs_2014[35][3]
    slice_3b = dfs_2014[35][1]
    slice_3b = [element.strip() for element in slice_3b]

    value_counts_dict_3 = {
        slice_3b[1]: int(slice_3[1]), slice_3b[2]: int(slice_3[2]), slice_3b[3]: int(slice_3[3]), slice_3b[4]: int(slice_3[4]),
        slice_3b[3]: int(slice_3[3]), slice_3b[6]: int(slice_3[6]), slice_3b[3]: int(slice_3[3]), slice_3b[8]: int(slice_3[8])
    }

    slice_4 = dfs_2014[37][3]
    slice_4b = dfs_2014[37][1]
    slice_4b = [element.strip() for element in slice_4b]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3]), slice_4b[4]: int(slice_4[4]),
        slice_4b[5]: int(slice_4[5]), slice_4b[6]: int(slice_4[6]), slice_4b[4]: int(slice_4[4]), slice_4b[8]: int(slice_4[8]),
        slice_4b[9]: int(slice_4[9]), slice_4b[10]: int(slice_4[10]), slice_4b[11]: int(slice_4[11]), slice_4b[12]: int(slice_4[12]),
        slice_4b[13]: int(slice_4[13]), slice_4b[14]: int(slice_4[14])
    }

    slice_5 = dfs_2014[40][3]
    slice_5b = dfs_2014[40][1]
    slice_5b = [element.strip() for element in slice_5b]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[5]: int(slice_5[5]),
        slice_5b[5]: int(slice_5[5]), slice_5b[6]: int(slice_5[6]), slice_5b[7]: int(slice_5[7])
    }

    value_counts_dict_6 = {
        'More than 20 years': 1, '10 to 20 years': 0, '2 years': 3, '3 years': 0,
        '1 year or less': 2, '5 to 10 years': 3, '4 years': 0
    }

    slice_2 = dfs_2014[38][3]
    slice_2b = dfs_2014[38][1]
    slice_2b = [element.strip() for element in slice_2b]

    value_counts_dict_2 = {
        slice_2b[1]: int(slice_2[1]), slice_2b[2]: int(slice_2[2]), slice_2b[3]: int(slice_2[3]), slice_2b[4]: int(slice_2[4])
    }

    slice_1 = dfs_2014[46][3]
    slice_1b = dfs_2014[46][1]
    slice_1b = [element.strip() for element in slice_1b]

    value_counts_dict_1 = {
        slice_1b[1]: int(slice_1[1]), slice_1b[2]: int(slice_1[2]), slice_1b[3]: int(slice_1[3]), slice_1b[4]: int(slice_1[4]), slice_1b[5]: int(slice_1[5])
    }

    print(value_counts_dict_1)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6 = dicts

    value_counts_dicts = {
        "Have you attended the SD conference before?": value_counts_dict_2,
        "What is your profession? Please select all that apply.": value_counts_dict_3,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_4,
        "What is your geographic region?": value_counts_dict_5,
        "If you have used the Conference Online Scheduler, which device have you used? Please check all that apply.": value_counts_dict_1,
        "How many years of experience do you have with system dynamics?": value_counts_dict_6,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2014' not in df.index:
        df.loc['2014'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2014', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2014', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # boolean dictionaries

    boolean_value_counts_dicts = {
        "An idea that is considered for future conferences, is to make the conference schedule available using a smartphone app. For that reason, we would like to ask you if you own a smartphone.": {"Yes": 117, "No": 18},
        "Have you used the Conference Online Scheduler?": {"Yes": 50, "No": 84},
    }

    print(boolean_value_counts_dicts)

    df = pd.read_csv("boolean_questions.csv", index_col=0)

    if '2014' not in df.index:
        df.loc['2014'] = pd.NA

    for col_name, col_data in boolean_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2014', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2014', col_name] = col_data

    df.to_csv("boolean_questions.csv")

    # numerical dictionaries

    print(dfs_2014[29])

    numerical_indices = [0, 2, 4, 6, 8, 10, 12, 14, 16, 18, 20, 22, 23, 24,
                         25, 26, 28, 29, 30]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        print(idx)
        slice_main = dfs_2014[idx][3]
        slice_helper = dfs_2014[idx][1]
        slice_helper = [element.strip() for element in slice_helper]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to conference geographical location my evaluation is:",
        "When it comes to the conference venue (building and facilities) my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "When it comes to the conference fee my evaluation is:",
        "When it comes to the overnight accommodation my evaluation is:",
        "When it comes to the printed abstract proceedings for use at the conference my evaluation is:",
        "When it comes to the satchel my evaluation is:",
        "When it comes to the lunch provided onsite my evaluation is:",
        "Are the plenary sessions too short or too long?",
        "Should we have more or fewer plenary sessions?",
        "Are the parallel sessions too short or too long?",
        "Should we have more or fewer parallel sessions?",
        "How do you evaluate the following sessions and workshops? [In-presence poster session]",
        "What is your opinion about the duration of the breaks during the conference?",
        "Would you like to see more or less academic work (as compared to applications)?",
        "How satisfied are you with the workshops?",
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2014' not in df.index:
        df.loc['2014'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2014', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2014', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    print(dfs_2014[39])

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "Additional comments 7",
        "Please indicate whether you used a hotel, hostel or AirBnB and the name of your accommodation if applicable. Please also list any additional comments you would like to make",
        "Additional comments 8",
        "Additional comments 9",
        "Additional comments 10",
        "If you have any comments on session formats, please list them here.",
        "If you could make changes to the format of the workshops, what would those changes be?",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify)",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "Any other comments you may have"
    ]
    text_list = []
    text_indices = [1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21,
                    27, 31, 32, 33, 34, 36, 39, 41, 47]

    for i in text_indices:
        text_list.append(dfs_2014[i][0][1:].values)

    text_value_counts_dicts = dict(zip(questions, text_list))
    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2014' not in df.index:
        df.loc['2014'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2014', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2014', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2015():
    dfs_2015 = []
    document_2015 = Document("data/2015 Report survey here.docx")
    # print out all columns for categorization

    # for paragraph in document_2015.paragraphs:
    #     line = paragraph.text
    #     if re.match(r'^\d+\.\s', line):
    #         # split at first occurrence of ' - ' and take the part after it
    #         line = line.split('.', 1)[1]
    #         print(line)

    for table in document_2015.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        df = df.drop_duplicates()
        df = df.reset_index(drop=True)
        dfs_2015.append(df)

    print(len(dfs_2015))
    print(type(dfs_2015))

    print(type(dfs_2015[0]))

    print("data_type: " + str(type(dfs_2015)))
    print("data_type item: " + str(type(dfs_2015[0])))
    # print(dfs_2015[0:7])

    # categorical dictionaries

    # manual counting for ugly text data

    # print(dfs_2015[39])

    slice_3 = dfs_2015[39][4]
    slice_3b = dfs_2015[39][1]
    value_counts_dict_3 = {
        slice_3b[1]: int(slice_3[1]), slice_3b[2]: int(slice_3[2]), slice_3b[3]: int(slice_3[3]), slice_3b[4]: int(slice_3[4]),
        slice_3b[3]: int(slice_3[3]), slice_3b[6]: int(slice_3[6]), slice_3b[3]: int(slice_3[3]), slice_3b[8]: int(slice_3[8])
    }

    slice_4 = dfs_2015[41][4]
    slice_4b = dfs_2015[41][1]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3]), slice_4b[4]: int(slice_4[4]),
        slice_4b[5]: int(slice_4[5]), slice_4b[6]: int(slice_4[6]), slice_4b[4]: int(slice_4[4]), slice_4b[8]: int(slice_4[8]),
        slice_4b[9]: int(slice_4[9]), slice_4b[10]: int(slice_4[10]), slice_4b[11]: int(slice_4[11]), slice_4b[12]: int(slice_4[12]),
        slice_4b[13]: int(slice_4[13]), slice_4b[14]: int(slice_4[14])
    }

    slice_5 = dfs_2015[44][4]
    slice_5b = dfs_2015[44][1]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[5]: int(slice_5[5]),
        slice_5b[5]: int(slice_5[5]), slice_5b[6]: int(slice_5[6]), slice_5b[7]: int(slice_5[7])
    }

    value_counts_dict_6 = {
        'More than 20 years': 8, '10 to 20 years': 12, '2 years': 3, '3 years': 2,
        '1 year or less': 13, '5 to 10 years': 11, '4 years': 2
    }

    slice_7 = dfs_2015[50][4]
    slice_7b = dfs_2015[50][1]

    value_counts_dict_7 = {
        slice_7b[1]: int(slice_7[1]), slice_7b[2]: int(slice_7[2]), slice_7b[3]: int(slice_7[3])
    }

    slice_2 = dfs_2015[42][4]
    slice_2b = dfs_2015[42][1]

    value_counts_dict_2 = {
        slice_2b[1]: int(slice_2[1]), slice_2b[2]: int(slice_2[2]), slice_2b[3]: int(slice_2[3]), slice_2b[4]: int(slice_2[4])
    }

    slice_9 = dfs_2015[52][4]
    slice_9b = dfs_2015[52][1]

    value_counts_dict_9 = {
        slice_9b[1]: int(slice_9[1]), slice_9b[2]: int(slice_9[2]), slice_9b[3]: int(slice_9[3])
    }

    slice_10 = dfs_2015[53][4]
    slice_10b = dfs_2015[53][1]

    value_counts_dict_10 = {
        slice_10b[1]: int(slice_10[1]), slice_10b[2]: int(slice_10[2]), slice_10b[3]: int(slice_10[3]), slice_10b[4]: int(slice_10[4]), slice_10b[5]: int(slice_10[5])
    }

    slice_8 = dfs_2015[55][4]
    slice_8b = dfs_2015[55][1]

    value_counts_dict_8 = {
        slice_8b[1]: int(slice_8[1]), slice_8b[2]: int(slice_8[2]), slice_8b[3]: int(slice_8[3])
    }

    # print(value_counts_dict_8)

    dicts = [value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8, value_counts_dict_9, value_counts_dict_10]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8, value_counts_dict_9, value_counts_dict_10 = dicts

    value_counts_dicts = {
        "Have you attended the SD conference before?": value_counts_dict_2,
        "What is your profession? Please select all that apply.": value_counts_dict_3,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_4,
        "What is your geographic region?": value_counts_dict_5,
        "Have you used the conference schedule smartphone app?": value_counts_dict_7,
        "Have you used the Conference Online Scheduler? The Conference Online Scheduler is available to everyone on the Web Portal Log In page; it shows the entire conference by day and hour.": value_counts_dict_9,
        "If you have used the Conference Online Scheduler, which device have you used? Please check all that apply.": value_counts_dict_10,
        "Have you used your scheduled conference sessions": value_counts_dict_8,
        "How many years of experience do you have with system dynamics?": value_counts_dict_6,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2015' not in df.index:
        df.loc['2015'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2015', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2015', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # numerical dictionaries

    print(dfs_2015[53])

    numerical_indices = [1, 3, 5, 7, 9, 11, 13, 15, 17, 20, 21, 23, 24, 25, 28,
                         29, 30, 51, 54, 56]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        print(idx)
        slice_main = dfs_2015[idx][4]
        slice_helper = dfs_2015[idx][1]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to conference geographical location my evaluation is:",
        "When it comes to the conference venue (building and facilities) my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "When it comes to the conference fee my evaluation is:",
        "When it comes to the overnight accommodation my evaluation is:",
        "When it comes to the web proceedings my evaluation is:",
        "Are the plenary sessions too short or too long?",
        "Should we have more or fewer plenary sessions?",
        "Are the parallel sessions too short or too long?",
        "Should we have more or fewer parallel sessions?",
        "How do you evaluate the following sessions and workshops? [In-presence poster session]",
        "What is your opinion about the duration of the breaks during the conference?",
        "Would you like to see more or less academic work (as compared to applications)?",
        "How satisfied are you with the workshops?",
        "If you have used the conference schedule smartphone app, how satisfied were you?",
        "If you have used the Conference Online Scheduler, how satisfied were you?",
        "If you have used Your scheduled conference sessions, how satisfied were you?",
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2015' not in df.index:
        df.loc['2015'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2015', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2015', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "Additional comments 7",
        "Please indicate whether you used a hotel, hostel or AirBnB and the name of your accommodation if applicable. Please also list any additional comments you would like to make",
        "Additional comments 8",
        "If you have any comments on session formats, please list them here.",
        "If you could make changes to the format of the workshops, what would those changes be?",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify)",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "If you have any suggestions for improvements in the conference schedule (online or printed), please list them here."
    ]
    # print(dfs_2015[45][0][1:].values)
    text_list = []
    text_indices = [(2,), (4,), (6,), (8, ), (10,), (12,), (14, ), (16, ),
                    (18,), (26,), (31,), (32, 33), (34,), (35,), (40,), (43,), (45,), (57,)]
    for tuple in text_indices:
        temp_list = []
        for i in tuple:
            temp_list.append(dfs_2015[i][0][1:].values)
        text_list.append(temp_list)

    text_value_counts_dicts = dict(zip(questions, text_list))

    text_value_counts_dicts = {key: [item for sublist in value for item in sublist]
                               for key, value in text_value_counts_dicts.items()}

    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2015' not in df.index:
        df.loc['2015'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2015', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2015', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2016():
    dfs_2016 = []
    document_2016 = Document("data/2016 Report survey.docx")
    # print out all columns for categorization

    # for paragraph in document_2016.paragraphs:
    #     line = paragraph.text
    #     if line.startswith('Q') and not line.startswith('Qu'):
    #         # split at first occurrence of ' - ' and take the part after it
    #         line = line.split(' - ', 1)[1]
    #         print(line)

    for table in document_2016.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2016.append(df)

    print(len(dfs_2016))
    print(type(dfs_2016))

    print(type(dfs_2016[0]))

    print("data_type: " + str(type(dfs_2016)))
    print("data_type item: " + str(type(dfs_2016[0])))
    # print(dfs_2016[0:7])

    # categorical dictionaries

    # manual counting for ugly text data

    print(dfs_2016[63])

    slice_3 = dfs_2016[63][2]
    slice_3b = dfs_2016[63][0]

    value_counts_dict_3 = {
        slice_3b[1]: int(slice_3[1]), slice_3b[2]: int(slice_3[2]), slice_3b[3]: int(slice_3[3]), slice_3b[4]: int(slice_3[4]),
        slice_3b[3]: int(slice_3[3]), slice_3b[6]: int(slice_3[6]), slice_3b[3]: int(slice_3[3]), slice_3b[8]: int(slice_3[8])
    }

    slice_4 = dfs_2016[66][2]
    slice_4b = dfs_2016[66][0]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3]), slice_4b[4]: int(slice_4[4]),
        slice_4b[5]: int(slice_4[5]), slice_4b[6]: int(slice_4[6]), slice_4b[4]: int(slice_4[4]), slice_4b[8]: int(slice_4[8]),
        slice_4b[9]: int(slice_4[9]), slice_4b[10]: int(slice_4[10]), slice_4b[11]: int(slice_4[11]), slice_4b[12]: int(slice_4[12]),
        slice_4b[13]: int(slice_4[13]), slice_4b[14]: int(slice_4[14])
    }

    slice_5 = dfs_2016[68][2]
    slice_5b = dfs_2016[68][0]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[5]: int(slice_5[5]),
        slice_5b[5]: int(slice_5[5]), slice_5b[6]: int(slice_5[6]), slice_5b[7]: int(slice_5[7])
    }

    value_counts_dict_6 = {
        'More than 20 years': 9, '10 to 20 years': 7, '2 years': 4, '3 years': 6,
        '1 year or less': 20, '5 to 10 years': 11, '4 years': 6
    }

    slice_7 = dfs_2016[77][2]
    slice_7b = dfs_2016[77][0]

    value_counts_dict_7 = {
        slice_7b[1]: int(slice_7[1]), slice_7b[2]: int(slice_7[2]), slice_7b[3]: int(slice_7[3])
    }

    slice_2 = dfs_2016[85][2]
    slice_2b = dfs_2016[85][0]

    value_counts_dict_2 = {
        slice_2b[1]: int(slice_2[1]), slice_2b[2]: int(slice_2[2]), slice_2b[3]: int(slice_2[3]), slice_2b[4]: int(slice_2[4])
    }

    slice_9 = dfs_2016[87][2]
    slice_9b = dfs_2016[87][0]

    value_counts_dict_9 = {
        slice_9b[1]: int(slice_9[1]), slice_9b[2]: int(slice_9[2]), slice_9b[3]: int(slice_9[3])
    }

    slice_10 = dfs_2016[89][2]
    slice_10b = dfs_2016[89][0]

    value_counts_dict_10 = {
        slice_10b[1]: int(slice_10[1]), slice_10b[2]: int(slice_10[2]), slice_10b[3]: int(slice_10[3]), slice_10b[4]: int(slice_10[4]), slice_10b[5]: int(slice_10[5])
    }

    slice_1 = dfs_2016[90][2]
    slice_1b = dfs_2016[90][0]

    value_counts_dict_1 = {
        slice_1b[1]: int(slice_1[1]), slice_1b[1]: int(slice_1[1]), slice_1b[3]: int(slice_1[3]), slice_1b[4]: int(slice_1[4])
    }

    slice_8 = dfs_2016[94][2]
    slice_8b = dfs_2016[94][0]

    value_counts_dict_8 = {
        slice_8b[1]: int(slice_8[1]), slice_8b[2]: int(slice_8[2]), slice_8b[3]: int(slice_8[3])
    }

    # print(value_counts_dict_8)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8, value_counts_dict_9, value_counts_dict_10]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8, value_counts_dict_9, value_counts_dict_10 = dicts

    value_counts_dicts = {
        "Have you attended the SD conference before?": value_counts_dict_2,
        "What is your profession? Please select all that apply.": value_counts_dict_3,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_4,
        "What is your geographic region?": value_counts_dict_5,
        "Have you used the conference schedule smartphone app?": value_counts_dict_7,
        "Have you used the Conference Online Scheduler? The Conference Online Scheduler is available to everyone on the Web Portal Log In page; it shows the entire conference by day and hour.": value_counts_dict_9,
        "If you have used the Conference Online Scheduler, which device have you used? Please check all that apply.": value_counts_dict_10,
        "Have you used your scheduled conference sessions": value_counts_dict_8,
        "Which type of poster session or sessions do you prefer?": value_counts_dict_1,
        "How many years of experience do you have with system dynamics?": value_counts_dict_6,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2016' not in df.index:
        df.loc['2016'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2016', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2016', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # numerical dictionaries

    print(dfs_2016[3])

    numerical_indices = [3, 7, 10, 13, 17, 20, 23, 27, 33, 35, 39, 41, 43, 48,
                         50, 79, 82, 92, 93, 95]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        print(idx)
        slice_main = dfs_2016[idx][2]
        slice_helper = dfs_2016[idx][0]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to conference geographical location my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "When it comes to the conference fee my evaluation is:",
        "When it comes to the overnight accommodation my evaluation is:",
        "When it comes to the web proceedings my evaluation is:",
        "Are the plenary sessions too short or too long?",
        "Should we have more or fewer plenary sessions?",
        "Are the parallel sessions too short or too long?",
        "Should we have more or fewer parallel sessions?",
        "How do you evaluate the following sessions and workshops? [In-presence poster session]",
        "Would you like to see more or less academic work (as compared to applications)?",
        "How satisfied are you with the workshops?",
        "When it comes to the conference venue (building and facilities) my evaluation is:",
        "What is your opinion about the duration of the breaks during the conference?",
        "If you have used the conference schedule smartphone app, how satisfied were you?",
        "If you have used the Conference Online Scheduler, how satisfied were you?",
        "If you have used Your scheduled conference sessions, how satisfied were you?",
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2016' not in df.index:
        df.loc['2016'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2016', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2016', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "Please indicate whether you used a hotel, hostel or AirBnB and the name of your accommodation if applicable. Please also list any additional comments you would like to make",
        "Additional comments 7",
        "If you have any comments on session formats, please list them here.",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify)",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "Additional comments 8",
        "If you could make changes to the format of the workshops, what would those changes be?",
        "Other (please specify) 4",
        "If you have any suggestions for improvements in the conference schedule (online or printed), please list them here."
    ]
    text_list = []
    text_indices = [(5, 6), (9,), (12,), (15, 16), (19,), (22,), (25, 26), (29, 30),
                    (45,), (52, 53, 54), (55, 56, 57), (58, 59, 60), (65,), (67,), (70,), (81,), (84,), (91,), (96,)]

    for tuple in text_indices:
        temp_list = []
        for i in tuple:
            temp_list.append(dfs_2016[i][0][1:].values)
        text_list.append(temp_list)

    text_value_counts_dicts = dict(zip(questions, text_list))

    text_value_counts_dicts = {key: [item for sublist in value for item in sublist]
                               for key, value in text_value_counts_dicts.items()}

    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2016' not in df.index:
        df.loc['2016'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2016', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2016', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2018():
    dfs_2018 = []
    document_2018 = Document("data/2018 ISDC survey report.docx")
    # print out all columns for categorization

    # for paragraph in document_2018.paragraphs:
    #     line = paragraph.text
    #     if line.startswith('Q'):
    #         # split at first occurrence of ' - ' and take the part after it
    #         line = line.split(' - ', 1)[1]
    #         print(line)

    for table in document_2018.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2018.append(df)

    print(len(dfs_2018))
    print(type(dfs_2018))

    print(type(dfs_2018[0]))

    print("data_type: " + str(type(dfs_2018)))
    print("data_type item: " + str(type(dfs_2018[0])))
    # print(dfs_2018[0:7])

    # categorical dictionaries

    # manual counting for ugly text data

    print(dfs_2018[70])

    slice_2 = (dfs_2018[70][3])

    value_counts_dict_2 = {
        'yes, 6 or more times': int(slice_2[4]), 'yes, 2-5 times': int(slice_2[3]), 'no, this was my first time attending the conference': int(slice_2[1]), 'yes, once before': int(slice_2[2])
    }

    slice_3 = dfs_2018[52][3]
    slice_3b = dfs_2018[52][1]

    value_counts_dict_3 = {
        slice_3b[1]: int(slice_3[1]), slice_3b[2]: int(slice_3[2]), slice_3b[3]: int(slice_3[3]), slice_3b[4]: int(slice_3[4]),
        slice_3b[3]: int(slice_3[3]), slice_3b[6]: int(slice_3[6]), slice_3b[3]: int(slice_3[3]), slice_3b[8]: int(slice_3[8])
    }

    slice_4 = dfs_2018[55][3]
    slice_4b = dfs_2018[55][1]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3]), slice_4b[4]: int(slice_4[4]),
        slice_4b[5]: int(slice_4[5]), slice_4b[6]: int(slice_4[6]), slice_4b[4]: int(slice_4[4]), slice_4b[8]: int(slice_4[8]),
        slice_4b[9]: int(slice_4[9]), slice_4b[10]: int(slice_4[10]), slice_4b[11]: int(slice_4[11]), slice_4b[12]: int(slice_4[12]),
        slice_4b[13]: int(slice_4[13]), slice_4b[14]: int(slice_4[14])
    }

    slice_5 = dfs_2018[57][3]
    slice_5b = dfs_2018[57][1]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[5]: int(slice_5[5]),
        slice_5b[5]: int(slice_5[5]), slice_5b[6]: int(slice_5[6]), slice_5b[7]: int(slice_5[7])
    }

    value_counts_dict_6 = {
        'More than 20 years': 10, '10 to 20 years': 17, '2 years': 9, '3 years': 6,
        '1 year or less': 18, '5 to 10 years': 12, '4 years': 2
    }

    slice_7 = dfs_2018[62][3]
    slice_7b = dfs_2018[62][1]

    value_counts_dict_7 = {
        slice_7b[1]: int(slice_7[1]), slice_7b[2]: int(slice_7[2]), slice_7b[3]: int(slice_7[3])
    }

    slice_1 = dfs_2018[72][3]
    slice_1b = dfs_2018[72][1]

    value_counts_dict_1 = {
        slice_1b[1]: int(slice_1[1]), slice_1b[2]: int(slice_1[2]), slice_1b[3]: int(slice_1[3]), slice_1b[4]: int(slice_1[4])
    }

    # print(value_counts_dict_1)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6, value_counts_dict_7]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6, value_counts_dict_7 = dicts

    value_counts_dicts = {
        "Have you attended the SD conference before?": value_counts_dict_2,
        "What is your profession? Please select all that apply.": value_counts_dict_3,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_4,
        "What is your geographic region?": value_counts_dict_5,
        "Have you used the conference schedule smartphone app?": value_counts_dict_7,
        "Which type of poster session or sessions do you prefer?": value_counts_dict_1,
        "How many years of experience do you have with system dynamics?": value_counts_dict_6,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2018' not in df.index:
        df.loc['2018'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2018', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2018', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # numerical dictionaries

    print(dfs_2018[72])

    numerical_indices = [2, 5, 8, 11, 14, 17, 20, 23, 28, 30, 34, 36, 38, 43,
                         45, 64, 67, 74, 77, 79, 81, 83, 85, 87, 89]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        slice_main = dfs_2018[idx][3]
        slice_helper = dfs_2018[idx][1]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to conference geographical location my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "When it comes to the conference fee my evaluation is:",
        "When it comes to the overnight accommodation my evaluation is:",
        "When it comes to the conference printed Program Schedule received at registration my evaluation is:",
        "Are the plenary sessions too short or too long?",
        "Should we have more or fewer plenary sessions?",
        "Are the parallel sessions too short or too long?",
        "Should we have more or fewer parallel sessions?",
        "How do you evaluate the following sessions and workshops? [In-presence poster session]",
        "Would you like to see more or less academic work (as compared to applications)?",
        "How satisfied are you with the workshops?",
        "When it comes to the conference venue (building and facilities) my evaluation is:",
        "What is your opinion about the duration of the breaks during the conference?",
        "If you have used the conference schedule smartphone app, how satisfied were you?",
        "When it comes to the conference Online Program and access to presented work my evaluation is:",
        "How satisfied are you with the education day poster sessions?",
        "How do you evaluate the following sessions and workshops? [Work-in-progress (WIP) sessions]",
        "To what extent do the work in progress (WIP) sessions add value to the conference program?",
        "How do you evaluate the following sessions and workshops? [Feedback sessions]",
        "To what extent do the feedback sessions add value to the conference program?",
        "How do you rate the overall quality of the presented work?",
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2018' not in df.index:
        df.loc['2018'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2018', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2018', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "Please indicate whether you used a hotel, hostel or AirBnB and the name of your accommodation if applicable. Please also list any additional comments you would like to make",
        "Additional comments 7",
        "If you have any comments on session formats, please list them here.",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify)",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "Additional comments 8",
        "If you could make changes to the format of the workshops, what would those changes be?",
        "Other (please specify) 4",
        "If you have any suggestions for improvements in the conference schedule (online or printed), please list them here."
    ]
    text_list = []
    text_indices = [4, 7, 10, 13, 16, 19, 22, 25,
                    40, 47, 48, 49, 54, 56, 59, 66, 69, 73, 75]

    for i in text_indices:
        text_list.append(dfs_2018[i][0][1:].values)

    text_value_counts_dicts = dict(zip(questions, text_list))
    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2018' not in df.index:
        df.loc['2018'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2018', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2018', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2019():
    dfs_2019 = []
    document_2019 = Document("data/2019 c ISDC.docx")
    # print out all columns for categorization

    # for paragraph in document_2019.paragraphs:
    #     line = paragraph.text
    #     if line.startswith('Q'):
    #         # split at first occurrence of ' - ' and take the part after it
    #         line = line.split(' - ', 1)[1]
    #         print(line)

    for table in document_2019.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2019.append(df)

    print(len(dfs_2019))
    print(type(dfs_2019))

    print(type(dfs_2019[0]))

    print("data_type: " + str(type(dfs_2019)))
    print("data_type item: " + str(type(dfs_2019[0])))
    # print(dfs_2019[0:7])

    # categorical dictionaries

    slice_1 = dfs_2019[54][3]
    slice_1b = dfs_2019[54][1]

    value_counts_dict_1 = {
        slice_1b[1]: int(slice_1[1]), slice_1b[2]: int(slice_1[2]), slice_1b[3]: int(slice_1[3]), slice_1b[4]: int(slice_1[4])
    }

    # # manual counting for ugly text data

    slice_2 = (dfs_2019[78][3])

    value_counts_dict_2 = {
        'yes, 6 or more times': int(slice_2[4]), 'yes, 2-5 times': int(slice_2[3]), 'no, this was my first time attending the conference': int(slice_2[1]), 'yes, once before': int(slice_2[2])
    }

    slice_3 = dfs_2019[74][3]
    slice_3b = dfs_2019[74][1]

    value_counts_dict_3 = {
        slice_3b[1]: int(slice_3[1]), slice_3b[2]: int(slice_3[2]), slice_3b[3]: int(slice_3[3]), slice_3b[4]: int(slice_3[4]),
        slice_3b[3]: int(slice_3[3]), slice_3b[6]: int(slice_3[6]), slice_3b[3]: int(slice_3[3]), slice_3b[8]: int(slice_3[8])
    }

    slice_4 = dfs_2019[76][3]
    slice_4b = dfs_2019[76][1]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3]), slice_4b[4]: int(slice_4[4]),
        slice_4b[5]: int(slice_4[5]), slice_4b[6]: int(slice_4[6]), slice_4b[4]: int(slice_4[4]), slice_4b[8]: int(slice_4[8]),
        slice_4b[9]: int(slice_4[9]), slice_4b[10]: int(slice_4[10]), slice_4b[11]: int(slice_4[11]), slice_4b[12]: int(slice_4[12]),
        slice_4b[13]: int(slice_4[13]), slice_4b[14]: int(slice_4[14])
    }

    slice_5 = dfs_2019[80][3]
    slice_5b = dfs_2019[80][1]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[5]: int(slice_5[5]),
        slice_5b[5]: int(slice_5[5]), slice_5b[6]: int(slice_5[6]), slice_5b[7]: int(slice_5[7])
    }

    value_counts_dict_6 = {
        'More than 20 years': 4, '10 to 20 years': 4, '2 years': 2, '3 years': 1,
        '1 year or less': 2, '5 to 10 years': 0, '4 years': 2
    }

    # print(value_counts_dict_5)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6 = dicts

    value_counts_dicts = {
        "Which type of poster session or sessions do you prefer?": value_counts_dict_1,
        "Have you attended the SD conference before?": value_counts_dict_2,
        "What is your profession? Please select all that apply.": value_counts_dict_3,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_4,
        "What is your geographic region?": value_counts_dict_5,
        "How many years of experience do you have with system dynamics?": value_counts_dict_6,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2019' not in df.index:
        df.loc['2019'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2019', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2019', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # numerical dictionaries

    print(dfs_2019[3])

    numerical_indices = [3, 6, 9, 12, 15, 18, 21, 24, 27, 29, 34, 36, 40,
                         42, 44, 46, 48, 50, 52, 58, 62, 64, 66]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        slice_main = dfs_2019[idx][3]
        slice_helper = dfs_2019[idx][1]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to conference geographical location my evaluation is:",
        "When it comes to the conference venue (building and facilities) my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "When it comes to the conference fee my evaluation is:",
        "When it comes to the overnight accommodation my evaluation is:",
        "When it comes to the conference printed Program Schedule received at registration my evaluation is:",
        "When it comes to the conference Online Program and access to presented work my evaluation is:",
        "Are the plenary sessions too short or too long?",
        "Should we have more or fewer plenary sessions?",
        "Are the parallel sessions too short or too long?",
        "Should we have more or fewer parallel sessions?",
        "How do you evaluate the following sessions and workshops? [In-presence poster session]",
        "How do you evaluate the following sessions and workshops? [Work-in-progress (WIP) sessions]",
        "To what extent do the work in progress (WIP) sessions add value to the conference program?",
        "How do you evaluate the following sessions and workshops? [Feedback sessions]",
        "To what extent do the feedback sessions add value to the conference program?",
        "How do you rate the overall quality of the presented work?",
        "What is your opinion about the duration of the breaks during the conference?",
        "Would you like to see more or less academic work (as compared to applications)?",
        "How satisfied are you with the workshops?",
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2019' not in df.index:
        df.loc['2019'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2019', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2019', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "Additional comments 7",
        "Please indicate whether you used a hotel, hostel or AirBnB and the name of your accommodation if applicable. Please also list any additional comments you would like to make",
        "Additional comments 8",
        "Other (please specify)",
        "If you have any comments on session formats, please list them here.",
        "If you could make changes to the format of the workshops, what would those changes be?",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "If you have any suggestions for improvements that we can make for future conferences, please list them here."
    ]
    text_list = []
    text_indices = [4, 7, 10, 13, 16, 19, 22, 25,
                    32, 55, 56, 67, 68, 69, 70, 75, 81, 84]

    for i in text_indices:
        text_list.append(dfs_2019[i][0][1:].values)

    text_value_counts_dicts = dict(zip(questions, text_list))
    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2019' not in df.index:
        df.loc['2019'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2019', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2019', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2020():
    dfs_2020 = []
    document_2020 = Document("data/2020 ISDC survey Report.docx")
    # print out all columns for categorization

    # for paragraph in document_2020.paragraphs:
    #     line = paragraph.text
    #     if line.startswith('Q'):
    #         # split at first occurrence of ' - ' and take the part after it
    #         line = line.split(' - ', 1)[1]
    #         print(line)

    for table in document_2020.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2020.append(df)

    print(len(dfs_2020))
    print(type(dfs_2020))

    print(type(dfs_2020[0]))

    print("data_type: " + str(type(dfs_2020)))
    print("data_type item: " + str(type(dfs_2020[0])))
    # print(dfs_2020[0:7])

    # categorical dictionaries

    slice_1 = dfs_2020[0][3]

    value_counts_dict_1 = {
        "Attend workshops": int(slice_1[1]), "Access session recordings": int(slice_1[2]), "View posters or attend poster session": int(slice_1[3]), "View 3D app for student colloquium": int(slice_1[4])}

    slice_2 = dfs_2020[40][3]
    value_counts_dict_30 = {
        "View posters on website and meet authors on zoom": int(slice_2[1]), "View posters and meet with authors in 3D app": int(slice_2[2]), "No Opinion": int(slice_2[3])
    }

    # manual counting for ugly text data

    value_counts_dict_2 = {
        'More than 20 years': 26, '10 to 20 years': 36, '2 years': 22, '3 years': 18,
        '1 year or less': 15, '5 to 10 years': 48, '4 years': 7
    }

    slice_3 = (dfs_2020[43][3])

    value_counts_dict_3 = {
        'yes, 6 or more times': int(slice_3[4]), 'yes, 2-5 times': int(slice_3[3]), 'no, this was my first time attending the conference': int(slice_3[1]), 'yes, once before': int(slice_3[2])
    }

    slice_4 = dfs_2020[45][3]
    slice_4b = dfs_2020[45][1]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3])
    }

    slice_5 = dfs_2020[46][3]
    slice_5b = dfs_2020[46][1]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[4]: int(slice_5[4]), slice_5b[5]: int(slice_5[5])
    }

    slice_6 = dfs_2020[55][3]
    slice_6b = dfs_2020[55][1]

    value_counts_dict_6 = {
        slice_6b[1]: int(slice_6[1]), slice_6b[2]: int(slice_6[2]), slice_6b[3]: int(slice_6[3]), slice_6b[4]: int(slice_6[4]),
        slice_6b[5]: int(slice_6[5]), slice_6b[6]: int(slice_6[6]), slice_6b[7]: int(slice_6[7]), slice_6b[8]: int(slice_6[8]),
        slice_6b[9]: int(slice_6[9]),
    }

    slice_7 = dfs_2020[57][3]
    slice_7b = dfs_2020[57][1]

    value_counts_dict_7 = {
        slice_7b[1]: int(slice_7[1]), slice_7b[2]: int(slice_7[2]), slice_7b[3]: int(slice_7[3]), slice_7b[4]: int(slice_7[4]),
        slice_7b[5]: int(slice_7[5]), slice_7b[6]: int(slice_7[6]), slice_7b[7]: int(slice_7[7]), slice_7b[8]: int(slice_7[8]),
        slice_7b[9]: int(slice_7[9]), slice_7b[10]: int(slice_7[10]), slice_7b[11]: int(slice_7[11]), slice_7b[12]: int(slice_7[12]),
        slice_7b[13]: int(slice_7[13]), slice_7b[14]: int(slice_7[14])
    }

    slice_8 = dfs_2020[60][3]
    slice_8b = dfs_2020[60][1]

    value_counts_dict_8 = {
        slice_8b[1]: int(slice_8[1]), slice_8b[2]: int(slice_8[2]), slice_8b[3]: int(slice_8[3]), slice_8b[4]: int(slice_8[4]),
        slice_8b[5]: int(slice_8[5]), slice_8b[6]: int(slice_8[6])
    }

    # print(value_counts_dict_8)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8,
             value_counts_dict_30]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8, value_counts_dict_30 = dicts

    value_counts_dicts = {
        "What aspects of the virtual conference did you experience? Select all that apply.": value_counts_dict_1,
        "For our virtual conference, student colloquium posters were displayed in the 3D app and other conference posters were displayed on the conference website. Please indicate which of the following formats you prefer (select one):": value_counts_dict_30,
        "How many years of experience do you have with system dynamics?": value_counts_dict_2,
        "Have you attended the SD conference before?": value_counts_dict_3,
        "How do you evaluate this year’s hybrid format as compared to the purely in presence and virtual formats?": value_counts_dict_4,
        "Which of the following features, if any, would be worthwhile to continue for future conferences? Please select all that apply.": value_counts_dict_5,
        "What is your profession? Please select all that apply.": value_counts_dict_6,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_7,
        "What is your geographic region?": value_counts_dict_8,
    }

    print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2020' not in df.index:
        df.loc['2020'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2020', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2020', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # boolean dictionaries

    indices = [52, 54]
    questions = [
        "Have you contacted any presenters for more information?",
        "Did you visit any exhibitors during the conference?",
    ]

    boolean_value_counts_dicts = {}

    for idx, question in zip(indices, questions):
        slice_data = dfs_2020[idx][3]
        value_counts_dict = {
            'Yes': int(slice_data[1]),
            'No': int(slice_data[2])
        }
        boolean_value_counts_dicts[question] = value_counts_dict

    print(boolean_value_counts_dicts)

    df = pd.read_csv("boolean_questions.csv", index_col=0)

    if '2020' not in df.index:
        df.loc['2020'] = pd.NA

    for col_name, col_data in boolean_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2020', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2020', col_name] = col_data

    df.to_csv("boolean_questions.csv")

    # numerical dictionaries

    numerical_indices = [2, 5, 8, 11, 14, 17,
                         20, 22, 24, 26, 28, 30, 32, 34, 37]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        slice_main = dfs_2020[idx][3]
        slice_helper = dfs_2020[idx][1]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the conference website and access to presented work, my evaluation is:",
        "When it comes to the 3D app, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "How do you evaluate the following sessions and workshops? [Plenary sessions]",
        "How do you evaluate the following sessions and workshops? [Parallel sessions]",
        "How do you evaluate the following sessions and workshops? [Work-in-progress (WIP) sessions]",
        "How do you evaluate the following sessions and workshops? [Dialog sessions]",
        "How do you evaluate the following sessions and workshops? [Feedback sessions]",
        "How do you evaluate the following sessions and workshops? [Student-Organized Colloquium]",
        "How do you evaluate the following sessions and workshops? [Virtual poster sessions]",
        "How do you evaluate the following sessions and workshops? [Online workshops]",
        "How do you rate the overall quality of the presented work?"
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2020' not in df.index:
        df.loc['2020'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2020', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2020', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "Additional comments 6",
        "If you have any comments on session formats, please list them here.",
        "Additional comments 7",
        "Other (please specify)",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "Other (please specify) 4",
        "If you have any suggestions for improvements that we can make for future conferences, please list them here."
    ]
    text_list = []
    text_indices = [3, 6, 9, 12, 15, 18, 35,
                    38, 47, 48, 49, 50, 56, 58, 61, 62]

    for i in text_indices:
        text_list.append(dfs_2020[i][0][1:].values)

    text_value_counts_dicts = dict(zip(questions, text_list))
    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2020' not in df.index:
        df.loc['2020'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2020', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2020', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2021():
    dfs_2021 = []
    document_2021 = Document("data/2021 b ISDC.docx")

    # print out all columns for categorization

    # for paragraph in document_2021.paragraphs:
    #     line = paragraph.text
    #     if line.startswith('Q'):
    #         # split at first occurrence of ' - ' and take the part after it
    #         line = line.split(' - ', 1)[1]
    #         print(line)

    for table in document_2021.tables:
        text = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(text)
        dfs_2021.append(df)

    print(len(dfs_2021))
    print(type(dfs_2021))

    print(type(dfs_2021[0]))

    print("data_type: " + str(type(dfs_2021)))
    print("data_type item: " + str(type(dfs_2021[0])))
    # print(dfs_2021[0:7])

    # categorical dictionaries

    slice_1 = dfs_2021[0][3]

    value_counts_dict_1 = {
        "Attend workshops": int(slice_1[1]), "Access session recordings": int(slice_1[2]), "View posters or attend poster session": int(slice_1[3])}

    # manual counting for ugly text data

    value_counts_dict_2 = {
        'More than 20 years': 15, '10 to 20 years': 11, '2 years': 9, '3 years': 9,
        '1 year or less': 4, '5 to 10 years': 4, '4 years': 2
    }

    slice_3 = (dfs_2021[38][3])

    value_counts_dict_3 = {
        'yes, 6 or more times': int(slice_3[4]), 'yes, 2-5 times': int(slice_3[3]), 'no, this was my first time attending the conference': int(slice_3[1]), 'yes, once before': int(slice_3[2])
    }

    slice_4 = dfs_2021[40][3]
    slice_4b = dfs_2021[40][1]

    value_counts_dict_4 = {
        slice_4b[1]: int(slice_4[1]), slice_4b[2]: int(slice_4[2]), slice_4b[3]: int(slice_4[3])
    }

    slice_5 = dfs_2021[41][3]
    slice_5b = dfs_2021[41][1]

    value_counts_dict_5 = {
        slice_5b[1]: int(slice_5[1]), slice_5b[2]: int(slice_5[2]), slice_5b[3]: int(slice_5[3]), slice_5b[4]: int(slice_5[4]), slice_5b[5]: int(slice_5[5])
    }

    slice_6 = dfs_2021[50][3]
    slice_6b = dfs_2021[50][1]

    value_counts_dict_6 = {
        slice_6b[1]: int(slice_6[1]), slice_6b[2]: int(slice_6[2]), slice_6b[3]: int(slice_6[3]), slice_6b[4]: int(slice_6[4]),
        slice_6b[5]: int(slice_6[5]), slice_6b[6]: int(slice_6[6]), slice_6b[7]: int(slice_6[7]), slice_6b[8]: int(slice_6[8]),
        slice_6b[9]: int(slice_6[9]),
    }

    slice_7 = dfs_2021[52][3]
    slice_7b = dfs_2021[52][1]

    value_counts_dict_7 = {
        slice_7b[1]: int(slice_7[1]), slice_7b[2]: int(slice_7[2]), slice_7b[3]: int(slice_7[3]), slice_7b[4]: int(slice_7[4]),
        slice_7b[5]: int(slice_7[5]), slice_7b[6]: int(slice_7[6]), slice_7b[7]: int(slice_7[7]), slice_7b[8]: int(slice_7[8]),
        slice_7b[9]: int(slice_7[9]), slice_7b[10]: int(slice_7[10]), slice_7b[11]: int(slice_7[11]), slice_7b[12]: int(slice_7[12]),
        slice_7b[13]: int(slice_7[13]), slice_7b[14]: int(slice_7[14])
    }

    slice_8 = dfs_2021[55][3]
    slice_8b = dfs_2021[55][1]

    value_counts_dict_8 = {
        slice_8b[1]: int(slice_8[1]), slice_8b[2]: int(slice_8[2]), slice_8b[3]: int(slice_8[3]), slice_8b[4]: int(slice_8[4]),
        slice_8b[5]: int(slice_8[5]), slice_8b[6]: int(slice_8[6])
    }

    slice_9 = dfs_2021[63][3]
    slice_9b = dfs_2021[63][1]

    value_counts_dict_9 = {
        slice_9b[1]: int(slice_9[1]), slice_9b[2]: int(slice_9[2]), slice_9b[3]: int(slice_9[3])
    }

    print(value_counts_dict_9)

    dicts = [value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4,
             value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8,
             value_counts_dict_9]

    for i, d in enumerate(dicts):
        dicts[i] = {k: v for k, v in sorted(
            d.items(), key=lambda item: item[1], reverse=True)}

    value_counts_dict_1, value_counts_dict_2, value_counts_dict_3, value_counts_dict_4, value_counts_dict_5, value_counts_dict_6, value_counts_dict_7, value_counts_dict_8, value_counts_dict_9 = dicts

    value_counts_dicts = {
        "What aspects of the virtual conference did you experience? Select all that apply.": value_counts_dict_1,
        "How many years of experience do you have with system dynamics?": value_counts_dict_2,
        "Have you attended the SD conference before?": value_counts_dict_3,
        "How do you evaluate this year’s hybrid format as compared to the purely in presence and virtual formats?": value_counts_dict_4,
        "Which of the following features, if any, would be worthwhile to continue for future conferences? Please select all that apply.": value_counts_dict_5,
        "What is your profession? Please select all that apply.": value_counts_dict_6,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_7,
        "What is your geographic region?": value_counts_dict_8,
        "The next conference is likely to be a hybrid (virtual and in-person) event. In which way would you prefer to attend?": value_counts_dict_9,
    }

    # print(value_counts_dicts)

    df = pd.read_csv("categorical_questions.csv", index_col=0)

    if '2021' not in df.index:
        df.loc['2021'] = pd.NA

    for col_name, col_data in value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2021', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2021', col_name] = col_data

    df.to_csv("categorical_questions.csv")

    # boolean dictionaries

    indices = [47, 49, 59, 61]
    questions = [
        "Have you contacted any presenters for more information?",
        "Did you visit any exhibitors during the conference?",
        "Would you recommend this event to others?",
        "Do you plan to attend this conference in the future?"
    ]

    boolean_value_counts_dicts = {}

    for idx, question in zip(indices, questions):
        slice_data = dfs_2021[idx][3]
        value_counts_dict = {
            'Yes': int(slice_data[1]),
            'No': int(slice_data[2])
        }
        boolean_value_counts_dicts[question] = value_counts_dict

    print(boolean_value_counts_dicts)

    df = pd.read_csv("boolean_questions.csv", index_col=0)

    if '2021' not in df.index:
        df.loc['2021'] = pd.NA

    for col_name, col_data in boolean_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2021', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2021', col_name] = col_data

    df.to_csv("boolean_questions.csv")

    # numerical dictionaries

    numerical_indices = [2, 5, 8, 11, 14, 17, 19, 21, 23, 25, 27, 29, 31, 34]
    value_counts_dict_list = []

    # Iterate over special indices
    for idx in numerical_indices:
        slice_main = dfs_2021[idx][3]
        slice_helper = dfs_2021[idx][1]
        value_counts_dict = {
            slice_helper[i]: int(slice_main[i]) for i in range(1, 8)
        }
        sorted_value_counts_dict = dict(
            sorted(value_counts_dict.items(), key=lambda item: item[1], reverse=True))
        value_counts_dict_list.append(sorted_value_counts_dict)

    # Map each dictionary to its respective question
    questions = [
        "When it comes to the content of the conference program, my evaluation is:",
        "When it comes to the conference website and access to presented work, my evaluation is:",
        "When it comes to the services provided by the conference organization, including technical support, my evaluation is:",
        "When it comes to the opportunity to socialize at the conference, my evaluation is:",
        "When it comes to overall conference value, my evaluation is:",
        "How do you evaluate the following sessions and workshops? [Plenary sessions]",
        "How do you evaluate the following sessions and workshops? [Parallel sessions]",
        "How do you evaluate the following sessions and workshops? [Work-in-progress (WIP) sessions]",
        "How do you evaluate the following sessions and workshops? [Dialog sessions]",
        "How do you evaluate the following sessions and workshops? [Feedback sessions]",
        "How do you evaluate the following sessions and workshops? [Student-Organized Colloquium]",
        "How do you evaluate the following sessions and workshops? [Virtual poster sessions]",
        "How do you evaluate the following sessions and workshops? [Online workshops]",
        "How do you rate the overall quality of the presented work?"
    ]

    numerical_value_counts_dicts = dict(zip(questions, value_counts_dict_list))

    print(numerical_value_counts_dicts)

    df = pd.read_csv("numerical_questions.csv", index_col=0)

    if '2021' not in df.index:
        df.loc['2021'] = pd.NA

    for col_name, col_data in numerical_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2021', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2021', col_name] = col_data

    df.to_csv("numerical_questions.csv")

    # text dictionaries

    questions = [
        "Additional comments",
        "Additional comments 2",
        "Additional comments 3",
        "Additional comments 4",
        "Additional comments 5",
        "If you have any comments on session formats, please list them here.",
        "Additional comments 6",
        "Other (please specify)",
        "What do you value most about participating in the conference?",
        "What was the best thing that happened to you at the conference?",
        "What was the worst thing that happened to you at the conference?",
        "Other (please specify) 2",
        "Other (please specify) 3",
        "Other (please specify) 4",
        "If you have any suggestions for improvements that we can make for future conferences, please list them here."
    ]
    text_list = []
    text_indices = [3, 6, 9, 12, 15, 32, 35, 42, 43, 44, 45, 51, 53, 56, 57]

    for i in text_indices:
        text_list.append(dfs_2021[i][0][1:].values)

    text_value_counts_dicts = dict(zip(questions, text_list))
    print(text_value_counts_dicts)

    df = pd.read_csv("text_responses.csv", index_col=0)

    if '2021' not in df.index:
        df.loc['2021'] = pd.NA

    for col_name, col_data in text_value_counts_dicts.items():
        if col_name in df.columns:
            df.at['2021', col_name] = col_data  # assuming year is the index
        else:
            # Add a new column with NaN for all previous rows
            df[col_name] = pd.NA
            df.at['2021', col_name] = col_data

    df.to_csv("text_responses.csv")


def year_2022():
    file_path = "data/2022 conference survey.xlsx"
    df_2022 = pd.read_excel(file_path)
    df_2022 = df_2022.drop("Timestamp", axis=1)
    pd.set_option('display.max_columns', None)
    print(df_2022.head())
    subset_indices = [0, 1, 25, 26, 27, 28, 29, 35, 36,
                      37, 38, 39, 40]  # Assuming zero-based indexing
    categorical_df_2022 = df_2022.iloc[:, subset_indices]
    print(categorical_df_2022.columns)

    print(
        categorical_df_2022["What is your geographic region?"].value_counts())

    value_counts_dict_1 = categorical_df_2022["How did you attend ISDC 2022?"].value_counts(
    ).to_dict()

    options_for_website_use = [
        'Attend zoom sessions (parallels, WIPs, roundtables, etc)',
        'Attend workshops',
        'Access session chat',
        'View posters or attend poster session',
        'Access session recordings'
    ]

    value_counts_dict_2 = explode_for_multiselect(
        categorical_df_2022, "What did you use the conference WEBSITE for? Please select all that apply.", options_for_website_use)

    value_counts_dict_3 = categorical_df_2022["How many years of experience do you have with system dynamics?"].value_counts(
    ).to_dict()

    value_counts_dict_4 = categorical_df_2022["Have you attended the SD conference before?"].value_counts(
    ).to_dict()

    value_counts_dict_5 = categorical_df_2022["If you attended a conference before, which formats did you experience?"].value_counts(
    ).to_dict()

    value_counts_dict_6 = categorical_df_2022["How do you evaluate this year’s hybrid format as compared to the purely in presence and virtual formats?"].value_counts(
    ).to_dict()

    options_for_future_conferences = [
        'access to recorded sessions',
        'networking in Zoom',
        'availability of pre-recorded talks'
    ]
    value_counts_dict_7 = explode_for_multiselect(
        categorical_df_2022, "Which of the following features, if any, would be worthwhile to continue for future conferences? Please select all that apply.", options_for_future_conferences)

    options_for_profession = [
        'Consulting',
        'Higher education',
        'In-house SD practitioner (private sector)',
        'In-house SD practitioner (public sector)',
        'K-12 education',
        'Research using SD',
        'SD client/customer',
        'Student',
        'Retired',
    ]

    value_counts_dict_8 = explode_for_multiselect(
        categorical_df_2022, "What is your profession? Please select all that apply.", options_for_profession)

    options_for_interest = [
        'Business policy',
        'Teaching',
        'Strategy',
        'Health',
        'Economic dynamics',
        'Energy and resources',
        'Environment and ecology',
        'Information science',
        'Methodology',
        'Operations management and supply chains',
        'Participatory problem solving',
        'Public policy',
        'Security',
        'Social and organizational dynamics',
    ]

    value_counts_dict_9 = explode_for_multiselect(
        categorical_df_2022, "What are your fields of interest? Please select all that apply.", options_for_interest)

    value_counts_dict_10 = categorical_df_2022["What is your geographic region?"].value_counts(
    ).to_dict()

    value_counts_dict_11 = categorical_df_2022["How do you self-identify in terms of gender?"].value_counts(
    ).to_dict()

    value_counts_dict_12 = categorical_df_2022["How old are you?"].value_counts(
    ).to_dict()

    options_for_values = [
        'Location',
        'Topic: System Dynamics',
        'Live presentations',
        'Recordings',
        'Social Interaction',
        'Q & A',
    ]

    value_counts_dict_13 = explode_for_multiselect(
        categorical_df_2022, "Which aspects do you value most when choosing to attend a conference? Select up to 3.", options_for_values)

    # create new DataFrame
    count_categorical_df_2022 = pd.DataFrame(index=["2022"])

    # store the dictionaries as single values in the DataFrame
    value_counts_dicts = {
        "How did you attend ISDC 2022?": value_counts_dict_1,
        "What did you use the conference WEBSITE for? Please select all that apply.": value_counts_dict_2,
        "How many years of experience do you have with system dynamics?": value_counts_dict_3,
        "Have you attended the SD conference before?": value_counts_dict_4,
        "If you attended a conference before, which formats did you experience?": value_counts_dict_5,
        "How do you evaluate this year’s hybrid format as compared to the purely in presence and virtual formats?": value_counts_dict_6,
        "Which of the following features, if any, would be worthwhile to continue for future conferences? Please select all that apply.": value_counts_dict_7,
        "What is your profession? Please select all that apply.": value_counts_dict_8,
        "What are your fields of interest? Please select all that apply.": value_counts_dict_9,
        "What is your geographic region?": value_counts_dict_10,
        "How do you self-identify in terms of gender?": value_counts_dict_11,
        "How old are you?": value_counts_dict_12,
        "Which aspects do you value most when choosing to attend a conference? Select up to 3.": value_counts_dict_13
    }

    for column_name, value_counts_dict in value_counts_dicts.items():
        count_categorical_df_2022.loc["2022",
                                      column_name] = [value_counts_dict]
    print(count_categorical_df_2022.shape)

    count_categorical_df_2022.to_csv("categorical_questions.csv")

    subset_indices = [33, 34, 42, 43]  # Assuming zero-based indexing
    boolean_df_2022 = df_2022.iloc[:, subset_indices]

    # print(boolean_df_2022.head())

    column_names = ["Have you contacted any presenters for more information?",
                    "Did you visit any exhibitors during the conference?",
                    "Would you recommend this event to others?",
                    "Do you plan to attend this conference in the future?"]

    count_boolean_df_2022 = pd.DataFrame(index=["2022"])

    for column_name in column_names:
        bool_counts = boolean_df_2022[column_name].value_counts().to_dict()
        count_boolean_df_2022.loc["2022", column_name] = [bool_counts]

    print(count_boolean_df_2022.shape)
    # print(count_boolean_df_2022)

    count_boolean_df_2022.to_csv("boolean_questions.csv")

    numerical_indices = [2, 4, 6, 8, 10, 12, 13, 14,
                         15, 16, 17, 18, 19, 20, 21, 23]  # Assuming zero-based indexing
    numerical_df_2022 = df_2022.iloc[:, numerical_indices]

    print(numerical_df_2022.columns)
    count_numerical_df_2022 = pd.DataFrame(index=["2022"])

    column_names = ["When it comes to the content of the conference program, my evaluation is:",
                    'When it comes to the conference website and access to presented work, my evaluation is:',
                    'When it comes to the services provided by the conference organization, including technical support, my evaluation is:',
                    'When it comes to the opportunity to socialize at the conference, my evaluation is:',
                    'When it comes to overall conference value, my evaluation is:',
                    'How do you evaluate the following sessions and workshops? [Plenary sessions]',
                    'How do you evaluate the following sessions and workshops? [Parallel sessions]',
                    'How do you evaluate the following sessions and workshops? [Work-in-progress (WIP) sessions]',
                    'How do you evaluate the following sessions and workshops? [Feedback sessions]',
                    'How do you evaluate the following sessions and workshops? [Student-Organized Colloquium on Monday]',
                    'How do you evaluate the following sessions and workshops? [Virtual poster sessions on Wednesday]',
                    'How do you evaluate the following sessions and workshops? [In-presence poster session on Wednessday]',
                    'How do you evaluate the following sessions and workshops? [Roundtables on Friday]',
                    'How do you evaluate the following sessions and workshops? [Online workshops on July 12]',
                    'How do you evaluate the following sessions and workshops? [Hybrid workshops on July 22]',
                    'How do you rate the overall quality of the presented work?']

    for column_name in column_names:
        bool_counts = numerical_df_2022[column_name].value_counts(
        ).to_dict()
        count_numerical_df_2022.loc["2022", column_name] = [bool_counts]

    print(count_numerical_df_2022.shape)
    # print(count_numerical_df_2022)

    count_numerical_df_2022.to_csv("numerical_questions.csv")

    text_indices = [3, 5, 7, 9, 11, 22, 24, 30, 31, 32, 41]
    text_df_2022 = df_2022.iloc[:, text_indices]

    print(text_df_2022.columns)
    # print(text_df_2022["Additional comments"])

    text_responses_2022 = pd.DataFrame(
        index=["2022"], columns=text_df_2022.columns)

    for column_name in text_df_2022.columns:
        # Drop NaN values and convert column to a list
        responses = text_df_2022[column_name].dropna().tolist()
        responses = np.array(responses)
        text_responses_2022.at["2022", column_name] = responses

    print(text_responses_2022.shape)
    print(text_responses_2022)

    text_responses_2022.to_csv("text_responses.csv")


def main():

    # RUN ALL 9 of them to generate the csv of the processed data!

    # year_2022()
    # year_2021()
    # year_2020()
    # year_2019()
    # year_2018()
    # year_2016()
    # year_2015()
    # year_2014()
    # year_2013()

    # Make big csv from pandas dataframe with year, categorical variables, Q1, Q2, Q19-23, Q26-31, 13 columns
    # Make big csv from pandas dataframe with year, boolean variables, Q24-25, Q32-33, 4 columns
    # Make big csv from pandas dataframe with year, numerical ratings, Q3-18, 16 columns
    # Make big csv from pandas dataframe with year, additional response strings, Addition Qs 1-5, others 11 columns
    pass


if __name__ == "__main__":
    main()
